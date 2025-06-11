mport os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# 학습 파일 목록 불러오기
def list_study_files():
    folder = "study_files"
    files = [f for f in os.listdir(folder) if f.endswith(".txt") or f.endswith(".md")]
    return files

# 학습 파일에서 퀴즈 문장 추출
def load_quiz_items(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        lines = f.readlines()
    return [line.strip() for line in lines if line.strip()]

# Word로 저장
def save_to_docx(quiz_items, filename="quiz_output.docx"):
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)
    full_path = os.path.join(output_dir, filename)

    doc = Document()

    # 한글 깨짐 방지: 글꼴 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    doc.add_heading('O/X 퀴즈 목록', level=1)
    for i, item in enumerate(quiz_items, 1):
        cleaned = item.replace("\n", " ")
        doc.add_paragraph(f"{i}. {cleaned}")
    doc.save(full_path)
    print(f"\n퀴즈가 Word 파일로 저장되었습니다: {full_path}\n")

# 메인 실행
def main():
    print("학습 파일 목록:")
    files = list_study_files()
    for i, f in enumerate(files, 1):
        print(f"{i}. {f}")

    try:
        choice = int(input("\n번호를 선택하세요: "))
        selected_file = files[choice - 1]
    except:
        print("잘못된 입력입니다.")
        return

    path = os.path.join("study_files", selected_file)
    quiz_items = load_quiz_items(path)

    print("\nO/X 퀴즈 출력:\n")
    for i, q in enumerate(quiz_items, 1):
        print(f"{i}. {q}")

    label = input("\n저장 라벨을 입력하세요 (예: 스크랩, 오답): ").strip()
    base_filename = os.path.splitext(selected_file)[0]
    filename = f"{base_filename}_{label}.docx"
    save_to_docx(quiz_items, filename)

if __name__ == "__main__":
    main()
