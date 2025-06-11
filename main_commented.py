import os
import datetime
from docx import Document  # Word 문서를 생성하고 저장하기 위한 모듈
from docx.shared import Pt  # 글자 크기를 포인트 단위로 설정하는 클래스
from docx.oxml.ns import qn  # 한글 글꼴 설정을 위한 네임스페이스 도우미

# --- 파일을 읽어서 문자열로 반환하는 함수 ---
def load_file(filepath):
    """
    주어진 경로의 파일을 열어 내용을 문자열로 반환한다.
    파일이 존재하지 않거나 인코딩 오류가 발생할 경우 None을 반환한다.
    """
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        print(f"[오류] 파일을 찾을 수 없습니다: {filepath}")
        return None

# --- 지정된 폴더 내 파일 목록을 출력하고 리스트로 반환하는 함수 ---
def list_files(folder):
    """
    주어진 폴더 내의 모든 파일을 리스트로 반환하며,
    사용자에게 번호와 함께 출력한다.
    """
    try:
        files = [f for f in os.listdir(folder) if os.path.isfile(os.path.join(folder, f))]
        if not files:
            print(f"[안내] '{folder}' 폴더에 파일이 없습니다.")
            return []
        print("\n선택 가능한 파일 목록:")
        for i, f in enumerate(files, 1):
            print(f"{i}. {f}")
        return files
    except FileNotFoundError:
        print(f"[오류] 폴더 '{folder}'를 찾을 수 없습니다.")
        return []

# --- 텍스트를 문장 단위로 나누고 각 문장에 O/X 항목을 추가하는 함수 ---
def generate_ox_quiz(text):
    """
    텍스트를 온점(.)을 기준으로 분리한 후, 각 문장 끝에 '(O/X)'를 붙인다.
    빈 문장은 제외한다.
    """
    sentences = [s.strip() for s in text.split('.') if s.strip()]
    return [s + '. (O/X)' for s in sentences]

# --- 퀴즈 문장을 Word 문서로 저장하는 함수 ---
def save_to_docx(quiz_items, filename="quiz_output.docx"):
    """
    퀴즈 문장 리스트를 받아 Word 문서를 생성하고,
    output 폴더에 지정한 파일명으로 저장한다.
    """
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)  # output 폴더 없으면 생성
    full_path = os.path.join(output_dir, filename)

    doc = Document()

    # Word 문서 스타일 설정 (맑은 고딕, 11pt)
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)
    if not style._element.rPr:
        from docx.oxml import OxmlElement
        style._element.rPr = OxmlElement('w:rPr')
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    doc.add_heading('O/X 퀴즈 목록', level=1)  # 제목 추가

    for i, item in enumerate(quiz_items, 1):
        cleaned_item = item.replace('\n', ' ')  # 줄바꿈 제거
        doc.add_paragraph(f"{i}. {cleaned_item}")  # 번호와 문장 추가

    doc.save(full_path)
    print(f"\n퀴즈가 Word 파일로 저장되었습니다: {full_path}")

# --- 사용자로부터 Y/N 입력을 받는 함수 ---
def ask_yes_no(prompt):
    """
    Y 또는 N을 입력할 때까지 계속 묻고,
    'y'는 True, 'n'은 False를 반환한다.
    """
    while True:
        ans = input(prompt).strip().lower()
        if ans == 'y':
            return True
        elif ans == 'n':
            return False
        else:
            print("[경고] Y 또는 N만 입력해주세요.")

# --- 메인 프로그램 흐름 ---
def main():
    folder = "study_files"  # 학습 파일이 있는 폴더명

    while True:
        files = list_files(folder)  # 파일 목록 출력 및 반환
        if not files:
            print("프로그램을 종료합니다.")
            break

        print("\n0. 종료하기")

        # 사용자로부터 번호 입력 받기
        try:
            choice = int(input("불러올 파일 번호를 입력하세요 (0 입력 시 종료): "))
            if choice == 0:
                print("프로그램을 종료합니다.")
                break
            choice -= 1
            if choice < 0 or choice >= len(files):
                print("[경고] 잘못된 번호입니다.\n")
                continue
        except ValueError:
            print("[경고] 숫자를 입력해야 합니다.\n")
            continue

        filename = files[choice]
        filepath = os.path.join(folder, filename)

        # 파일 확장자 검사
        if not filename.endswith(('.txt', '.md')):
            print("[경고] 지원하지 않는 파일 형식입니다.\n")
            continue

        text = load_file(filepath)
        if text:
            quiz_items = generate_ox_quiz(text)
            print("\nO/X 퀴즈 목록:\n")
            for i, q in enumerate(quiz_items, 1):
                print(f"{i}. {q}")

            # Word 저장 여부 묻기
            if ask_yes_no("\n현재 퀴즈를 Word(.docx) 파일로 저장하시겠습니까? (Y/N): "):
                docx_name = input("저장할 파일 이름을 입력하세요 (확장자 제외, 엔터 누르면 자동 생성): ").strip()
                if not docx_name:
                    now = datetime.datetime.now()
                    docx_name = now.strftime("quiz_%Y%m%d_%H%M%S")
                try:
                    save_to_docx(quiz_items, f"{docx_name}.docx")
                except Exception as e:
                    print(f"[오류] Word 저장에 실패했습니다: {e}")
        else:
            print("[오류] 파일을 불러오지 못했습니다.\n")
            continue

        # 반복 실행 여부 묻기
        if not ask_yes_no("\n프로그램을 다시 실행하시겠습니까? (Y/N): "):
            print("프로그램을 종료합니다. 감사합니다.")
            break

# --- 프로그램 시작 지점 ---
if __name__ == "__main__":
    main()
